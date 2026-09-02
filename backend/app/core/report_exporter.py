from __future__ import annotations

import re
from datetime import date, datetime
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape

from openpyxl import Workbook
from openpyxl.styles import Font
from openpyxl.utils import get_column_letter
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.dataset import Dataset
from app.models.datasource import DataSource
from app.models.report_template import ReportRun, ReportTemplate

# 依赖库：openpyxl(Excel)、python-docx(Word)、reportlab(PDF)；HTML 直接落盘无需额外依赖。
SUPPORTED_EXPORT_TYPES = {"excel", "pdf", "word", "html"}
EXPORT_ROW_LIMIT = 5000

_PLACEHOLDER_RE = re.compile(r"\{\{\s*([A-Za-z_][A-Za-z0-9_]*)\s*\}\}")
_SHEET_NAME_ILLEGAL_RE = re.compile(r"[\\/*?\[\]:]")

# AI 生成的报表 HTML 约定通过站内相对路径 /report-libs/echarts.min.js 引入 ECharts
# （前端静态资源）。下载后的 .html 用 file:// 打开时该路径无法加载，图表会空白，
# 因此导出时把运行库内联为自包含内容。
_ECHARTS_RELATIVE_SRC = "/report-libs/echarts.min.js"
_ECHARTS_CDN_SRC = "https://cdn.jsdelivr.net/npm/echarts@5/dist/echarts.min.js"
_ECHARTS_SCRIPT_TAG_RE = re.compile(
    r"<script\b[^>]*src=[\"']" + re.escape(_ECHARTS_RELATIVE_SRC) + r"[\"'][^>]*>\s*</script>",
    re.IGNORECASE | re.DOTALL,
)

_echarts_runtime: str | None = None
_echarts_runtime_missing = False

# 导出文件名统一用该时间戳，避免同一 run 内多次生成冲突。
_FILENAME_TS_FMT = "%Y%m%d%H%M%S%f"


def get_export_dir() -> Path:
    """解析导出文件目录（REPORT_EXPORT_DIR，默认 backend/exports），不存在则创建。"""
    raw = Path(settings.report_export_dir).expanduser()
    path = raw if raw.is_absolute() else Path(__file__).resolve().parents[2] / raw
    path.mkdir(parents=True, exist_ok=True)
    return path


def _render_text(value: Any, parameters: dict[str, Any]) -> Any:
    if not isinstance(value, str):
        return value

    def _replace(match: re.Match) -> str:
        key = match.group(1)
        if key in parameters:
            return str(parameters[key])
        return match.group(0)

    return _PLACEHOLDER_RE.sub(_replace, value)


def _safe_sheet_name(name: Any) -> str:
    cleaned = _SHEET_NAME_ILLEGAL_RE.sub(" ", str(name or "report")).strip() or "report"
    return cleaned[:31]


def _cell_value(value: Any) -> Any:
    if value is None or isinstance(value, (bool, int, float, str, datetime, date)):
        return value
    return str(value)


def _sanitize_text(value: Any) -> str:
    """把单元格值转为可写入 docx/pdf 的文本，剔除 XML 1.0 非法的控制字符。"""
    text = str(_cell_value(value))
    return "".join(
        ch for ch in text if ch in ("\t", "\n", "\r") or ord(ch) >= 32 or 0xE000 <= ord(ch) <= 0xF8FF
    )


def _pdf_paragraph_text(value: Any) -> str:
    """转义 XML 并把换行转为 <br/>，供 reportlab Paragraph 使用。"""
    return escape(_sanitize_text(value)).replace("\n", "<br/>")


def _layout_lines(template: ReportTemplate, parameters: dict[str, Any]) -> list[str]:
    """把 layout 的格子（支持 {{ 参数 }}）按行合并为文本行，供 PDF/Word 展示标题区。"""
    layout = template.layout_json if isinstance(template.layout_json, dict) else {}
    cells = layout.get("cells") if isinstance(layout.get("cells"), list) else []
    grid: dict[int, dict[int, str]] = {}
    for cell in cells:
        if not isinstance(cell, dict):
            continue
        try:
            row = int(cell.get("row") or 0)
            col = int(cell.get("col") or 0)
        except (TypeError, ValueError):
            continue
        text = str(_render_text(cell.get("value"), parameters) or "").strip()
        if text:
            grid.setdefault(row, {})[col] = text
    return ["  ".join(grid[row][col] for col in sorted(grid[row])) for row in sorted(grid)]


def _meta_text(template: ReportTemplate, dataset: Dataset, parameters: dict[str, Any]) -> str:
    parts = [
        f"版本 v{template.version}",
        f"数据集 {dataset.name}",
        f"导出时间 {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
    ]
    if parameters:
        parts.append("参数: " + ", ".join(f"{key}={value}" for key, value in parameters.items()))
    return " · ".join(parts)


def _output_path(export_dir: Path, run_id: int, suffix: str) -> Path:
    filename = f"report_run_{run_id}_{datetime.now().strftime(_FILENAME_TS_FMT)}{suffix}"
    return export_dir / filename


def build_excel_file(
    template: ReportTemplate,
    dataset: Dataset,
    columns: list[str],
    rows: list[dict[str, Any]],
    parameters: dict[str, Any],
    export_dir: Path,
    run_id: int,
) -> Path:
    """把模板 layout 单元格（支持 {{ 参数 }} 占位符）和数据集明细渲染成 xlsx。"""
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = _safe_sheet_name(template.name)

    sheet.cell(row=1, column=1, value=str(template.name)).font = Font(bold=True, size=14)
    sheet.cell(row=2, column=1, value=_meta_text(template, dataset, parameters))
    if parameters:
        sheet.cell(
            row=3,
            column=1,
            value="参数: " + ", ".join(f"{key}={value}" for key, value in parameters.items()),
        )

    row_cursor = 5
    layout = template.layout_json if isinstance(template.layout_json, dict) else {}
    cells = layout.get("cells") if isinstance(layout.get("cells"), list) else []
    max_layout_row = 0
    for cell in cells:
        if not isinstance(cell, dict):
            continue
        try:
            layout_row = int(cell.get("row") or 0)
            layout_col = int(cell.get("col") or 0)
        except (TypeError, ValueError):
            continue
        if layout_row < 1 or layout_col < 1:
            continue
        sheet.cell(
            row=row_cursor + layout_row - 1,
            column=layout_col,
            value=_render_text(cell.get("value"), parameters),
        )
        max_layout_row = max(max_layout_row, layout_row)
    if max_layout_row:
        row_cursor += max_layout_row + 1

    header_font = Font(bold=True)
    for index, column in enumerate(columns, start=1):
        sheet.cell(row=row_cursor, column=index, value=str(column)).font = header_font
    for row_offset, row in enumerate(rows, start=1):
        for column_index, column in enumerate(columns, start=1):
            sheet.cell(
                row=row_cursor + row_offset,
                column=column_index,
                value=_cell_value(row.get(column)),
            )
    for column_index, column in enumerate(columns, start=1):
        sheet.column_dimensions[get_column_letter(column_index)].width = max(
            12, min(40, len(str(column)) + 4)
        )

    path = _output_path(export_dir, run_id, ".xlsx")
    workbook.save(path)
    return path


def build_pdf_file(
    template: ReportTemplate,
    dataset: Dataset,
    columns: list[str],
    rows: list[dict[str, Any]],
    parameters: dict[str, Any],
    export_dir: Path,
    run_id: int,
) -> Path:
    """用 reportlab 把标题、layout 文本和数据集明细渲染成 PDF（内置 STSong-Light 支持中文）。"""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import LongTable, Paragraph, SimpleDocTemplate, Spacer, TableStyle

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    cn_font = "STSong-Light"

    title_style = ParagraphStyle("cn_title", fontName=cn_font, fontSize=17, leading=22, spaceAfter=8)
    meta_style = ParagraphStyle(
        "cn_meta",
        fontName=cn_font,
        fontSize=9,
        leading=14,
        spaceAfter=2,
        textColor=colors.HexColor("#637083"),
    )
    header_cell_style = ParagraphStyle(
        "cn_header_cell", fontName=cn_font, fontSize=9, leading=12, textColor=colors.white
    )
    cell_style = ParagraphStyle("cn_cell", fontName=cn_font, fontSize=8, leading=11)

    left_margin = right_margin = 16 * mm
    top_margin = bottom_margin = 14 * mm
    path = _output_path(export_dir, run_id, ".pdf")
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=left_margin,
        rightMargin=right_margin,
        topMargin=top_margin,
        bottomMargin=bottom_margin,
        title=str(template.name),
    )
    story: list[Any] = [
        Paragraph(_pdf_paragraph_text(template.name), title_style),
        Paragraph(_pdf_paragraph_text(_meta_text(template, dataset, parameters)), meta_style),
    ]
    layout_texts = _layout_lines(template, parameters)
    if layout_texts:
        story.append(Spacer(1, 2 * mm))
        for line in layout_texts:
            story.append(Paragraph(_pdf_paragraph_text(line), meta_style))
    story.append(Spacer(1, 4 * mm))

    if columns:
        data: list[list[Paragraph]] = [
            [Paragraph(escape(str(column)), header_cell_style) for column in columns]
        ]
        for row in rows:
            data.append([Paragraph(_pdf_paragraph_text(row.get(column)), cell_style) for column in columns])

        avail_width = A4[0] - left_margin - right_margin
        col_widths = [min(avail_width / len(columns), 56 * mm)] * len(columns)
        table = LongTable(data, repeatRows=1, colWidths=col_widths)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0f766e")),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d7e0ea")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f6f8fb")]),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]
            )
        )
        story.append(table)

    doc.build(story)
    return path


def build_word_file(
    template: ReportTemplate,
    dataset: Dataset,
    columns: list[str],
    rows: list[dict[str, Any]],
    parameters: dict[str, Any],
    export_dir: Path,
    run_id: int,
) -> Path:
    """用 python-docx 把标题、layout 文本和数据集明细渲染成 Word（.docx）。"""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    doc = Document()
    # 设置默认中文字体，避免打开文档时中文回退字体导致排版不一致。
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10)

    heading = doc.add_heading(level=0)
    run = heading.add_run(str(template.name))
    run.font.size = Pt(20)

    meta = doc.add_paragraph()
    meta_run = meta.add_run(_sanitize_text(_meta_text(template, dataset, parameters)))
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x63, 0x70, 0x83)

    for line in _layout_lines(template, parameters):
        doc.add_paragraph(_sanitize_text(line)).paragraph_format.space_after = Pt(2)

    if columns:
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        for index, column in enumerate(columns):
            header_cells[index].text = str(column)
        for row in rows:
            row_cells = table.add_row().cells
            for index, column in enumerate(columns):
                row_cells[index].text = _sanitize_text(row.get(column))

    path = _output_path(export_dir, run_id, ".docx")
    doc.save(path)
    return path


def _inline_echarts_runtime(html: str) -> str:
    """把 <script src="/report-libs/echarts.min.js"></script> 替换为内联脚本。

    后端 vendor 的 echarts.min.js 读取失败时回退为 CDN 地址，避免图表完全空白。
    未引用该站内路径的 HTML 原样返回。
    """
    global _echarts_runtime, _echarts_runtime_missing

    if not _ECHARTS_SCRIPT_TAG_RE.search(html):
        return html

    if _echarts_runtime is None and not _echarts_runtime_missing:
        vendor_path = Path(__file__).resolve().parents[1] / "static" / "report-libs" / "echarts.min.js"
        try:
            _echarts_runtime = vendor_path.read_text(encoding="utf-8")
        except OSError:
            _echarts_runtime_missing = True

    def _replace(_match: re.Match) -> str:
        if _echarts_runtime is not None:
            return f"<script>\n{_echarts_runtime}\n</script>"
        return f'<script src="{_ECHARTS_CDN_SRC}"></script>'

    return _ECHARTS_SCRIPT_TAG_RE.sub(_replace, html)


def build_html_file(template: ReportTemplate, export_dir: Path, run_id: int) -> Path:
    """把报表模板里的 HTML 内容（ai_html 等）落盘为可独立打开的 .html 文件。

    AI 生成的报表 HTML 通过站内相对路径引用 ECharts（仅在宿主页面内可加载），
    下载文件需要离线可打开，导出时会将运行库内联进 HTML。
    """
    layout = template.layout_json if isinstance(template.layout_json, dict) else {}
    html = layout.get("html")
    if not isinstance(html, str) or not html.strip():
        raise ValueError("该报表没有可下载的 HTML 内容")
    path = _output_path(export_dir, run_id, ".html")
    path.write_text(_inline_echarts_runtime(html), encoding="utf-8")
    return path


def _load_report_data(db: Session, template: ReportTemplate) -> tuple[Dataset, DataSource, list[str], list[dict[str, Any]]]:
    """查询模板关联数据集（列 + 行），Excel/PDF/Word 三格式共用一次取数。"""
    # 延迟导入，避免 core 模块在加载期依赖 api 层。
    from app.api.datasets import _execute_dataset_preview

    dataset = db.query(Dataset).filter(Dataset.id == template.dataset_id).first()
    if not dataset:
        raise ValueError("模板关联的数据集不存在")
    datasource = db.query(DataSource).filter(DataSource.id == dataset.datasource_id).first()
    if not datasource:
        raise ValueError("数据集关联的数据源不存在")

    result = _execute_dataset_preview(dataset, datasource, EXPORT_ROW_LIMIT)
    columns = [str(column) for column in (result.get("columns") or [])]
    rows = [dict(row) for row in (result.get("rows") or [])]
    return dataset, datasource, columns, rows


def execute_report_export(db: Session, template: ReportTemplate, run: ReportRun) -> tuple[Path, int]:
    """同步执行导出：按格式分派渲染，返回 (文件路径, 数据行数)。

    - html：无数据集报表（ai_html）的 layout html 直接落盘，行数为 0。
    - excel/pdf/word：查询模板关联数据集并渲染，需要模板已绑定数据集。
    """
    export_type = run.export_type
    if export_type not in SUPPORTED_EXPORT_TYPES:
        raise ValueError(f"暂不支持 {export_type} 导出，当前仅支持 {', '.join(sorted(SUPPORTED_EXPORT_TYPES))}")
    parameters = run.parameters_json if isinstance(run.parameters_json, dict) else {}
    export_dir = get_export_dir()

    if export_type == "html":
        path = build_html_file(template, export_dir, run.id)
        return path, 0

    dataset, _datasource, columns, rows = _load_report_data(db, template)
    if export_type == "excel":
        path = build_excel_file(template, dataset, columns, rows, parameters, export_dir, run.id)
    elif export_type == "pdf":
        path = build_pdf_file(template, dataset, columns, rows, parameters, export_dir, run.id)
    elif export_type == "word":
        path = build_word_file(template, dataset, columns, rows, parameters, export_dir, run.id)
    else:  # pragma: no cover - 已在上方校验
        raise ValueError(f"暂不支持 {export_type} 导出")
    return path, len(rows)
